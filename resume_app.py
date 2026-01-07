import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

st.set_page_config(
    page_title="高如慧 | Growth Marketing Lead",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 隱藏 Streamlit 預設元素
st.markdown("""
<style>
#MainMenu, footer, header, .stDeployButton {display: none !important;}
.stApp {background: #f5f5f5;}
.main .block-container {padding: 1rem; max-width: 900px;}
</style>
""", unsafe_allow_html=True)

# ==========================
# 嚴格比對後的資料變數 (供 HTML 與 Word 共用)
# ==========================
resume_data = {
    "name": "高如慧",
    "title": "Growth Marketing（AI-Driven）｜11 年跨產業行銷實戰",
    "summary": "以「結構化策略＋SEO＋成效型投放＋AI 自動化」建立可複製的成長系統。曾主導 POPRORO 0→1（8 個月達成月營收 800 萬）、推動 MacLove「蘋果二手」核心關鍵字取得 Google 首位，並以自動化流程讓代操時間下降約 66%。累積服務保健、美妝、3C、服飾、醫美等 10+ 品牌。",
    "contact": ["rhk9903@gmail.com", "0988-663-166", "新北市汐止區", "淡江大學 經濟學系"],
    "metrics": [
        ("800萬", "POPRORO 月營收（8 個月）"),
        ("ROAS 5", "保健食品投放模型"),
        ("#1", "蘋果二手（Google）"),
        ("-66%", "代操維運時間")
    ],
    "experience": [
        {
            "company": "個人接案｜行銷顧問",
            "date": "2025/6 ~ 現在", # 修正：PDF source: 30
            "role": "AI＋自動化成效系統",
            "desc": "服務服飾／美妝／保健／醫美等產業。將代操維運由 30 分鐘壓縮至約 10 分鐘。協助香港包包品牌達成年營收 +187%、ROAS 約 4。" # 修正：Doughnut -> 香港包包品牌
        },
        {
            "company": "森宏生技 LOVITA｜行銷專案經理",
            "date": "2025/5 ~ 2025/10",
            "role": "",
            "desc": "Non-Branding SEO（B 群）佔前 5 名中 3 名。建立穩定 ROAS 5。新品年度計畫透過 AI GTM 一週完成。"
        },
        {
            "company": "麥克愛愛 MacLove｜電商／行銷經理",
            "date": "2024/9 ~ 2025/3", # 修正：PDF source: 109 (2024/9)
            "role": "",
            "desc": "「蘋果二手」相關關鍵字取得 Google 第 1 名。Shopee 單月 89 萬（YoY +324%）。投放 ROAS 由 1 提升至 3。"
        },
        {
            "company": "歐賀服飾｜行銷經理",
            "date": "2023/12 ~ 2024/7",
            "role": "",
            "desc": "OMO 併入後電商營收 YoY +600%。LINE 會員 +700%，回購率 10%→25%。"
        },
        {
            "company": "高博士國際｜電商副理",
            "date": "2022/10 ~ 2023/11",
            "role": "",
            "desc": "打造「小白鞋」單月 490+ 雙。整體年營收 +30%。ROAS 維持 5+。擔任 CDP 導入與轉換流程改善。"
        },
        {
            "company": "個人接案｜行銷顧問", # 新增：補上 PDF source: 137 的經歷
            "date": "2022/4 ~ 2022/10",
            "role": "數位行銷顧問",
            "desc": "協助香港包包品牌、保健食品、醫美診所。保健品牌月營收由 100 萬穩定至 200 萬。"
        },
        {
            "company": "米波國際 meepShop｜行銷副理",
            "date": "2017/8 ~ 2022/3",
            "role": "",
            "desc": "主導 POPRORO 0→1，8 個月達成月營收 800 萬，淨利率 15–20%。管理 6 人團隊與月預算約 400 萬。"
        }
    ]
}

# ==========================
# 功能：生成 Word 檔
# ==========================
def create_word_resume():
    doc = Document()
    
    # 設定中文字型
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    
    # 標題 (姓名)
    h1 = doc.add_heading(resume_data['name'], 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 副標題
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(resume_data['title'])
    run_title.font.color.rgb = RGBColor(102, 102, 102)
    
    # 聯絡資訊
    p_contact = doc.add_paragraph(" | ".join(resume_data['contact']))
    p_contact.style = 'Quote'
    
    doc.add_paragraph("-" * 80)

    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(resume_data['summary'])
    
    # Metrics
    doc.add_heading('Key Metrics', level=1)
    table = doc.add_table(rows=1, cols=4)
    row = table.rows[0]
    for idx, (num, label) in enumerate(resume_data['metrics']):
        cell = row.cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(num + "\n")
        r1.bold = True
        r1.font.color.rgb = RGBColor(37, 99, 235) # Blue
        r2 = p.add_run(label)
        r2.font.size = Pt(9)

    # Experience
    doc.add_heading('Experience', level=1)
    
    for exp in resume_data['experience']:
        # 公司與日期一行
        p_exp = doc.add_paragraph()
        run_company = p_exp.add_run(exp['company'])
        run_company.bold = True
        run_company.font.size = Pt(11)
        
        run_date = p_exp.add_run(f"\t{exp['date']}")
        run_date.font.color.rgb = RGBColor(128, 128, 128)
        run_date.font.size = Pt(9)
        p_exp.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        
        # 描述
        if exp['role']:
            p_role = doc.add_paragraph()
            run_role = p_role.add_run(exp['role'])
            run_role.font.color.rgb = RGBColor(37, 99, 235)
            run_role.font.size = Pt(9)
        
        p_desc = doc.add_paragraph(exp['desc'])
        p_desc.paragraph_format.space_after = Pt(10)

    # Skills (Simplified for Word)
    doc.add_heading('Skills', level=1)
    p_skills = doc.add_paragraph()
    p_skills.add_run("Growth & Strategy: ").bold = True
    p_skills.add_run("0→1 品牌架構, 成長策略, LTV/CAC, 漏斗優化\n")
    
    p_skills.add_run("AI & Automation: ").bold = True
    p_skills.add_run("AI 內容產製, Streamlit Tooling, 自動化診斷\n")
    
    p_skills.add_run("Digital Marketing: ").bold = True
    p_skills.add_run("FB/Google Ads, SEO, LINE OA, 內容行銷\n")
    
    p_skills.add_run("Tools: ").bold = True
    p_skills.add_run("GA, GSC, Shopline, 91APP, Mixpanel")

    # 存入 BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================
# Sidebar: 下載功能
# ==========================
with st.sidebar:
    st.header("功能選單")
    docx_file = create_word_resume()
    st.download_button(
        label="📥 下載 Word 履歷",
        data=docx_file,
        file_name="高如慧_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ==========================
# HTML 版履歷 (更新後的內容)
# ==========================
# 動態生成經歷 HTML
exp_html_left = ""
exp_html_right = ""

# 將經歷拆分左右兩欄 (前3個左邊，後4個右邊)
for i, exp in enumerate(resume_data['experience']):
    item_html = f"""
    <div class="exp-item">
        <div class="exp-header">
            <span class="exp-company">{exp['company']}</span>
            <span class="exp-date">{exp['date']}</span>
        </div>
        {f'<div class="exp-role">{exp["role"]}</div>' if exp["role"] else ''}
        <div class="exp-desc">
            {exp['desc'].replace('800 萬', '<span class="hl">800 萬</span>')
                        .replace('ROAS 5', '<span class="hl">ROAS 5</span>')
                        .replace('+187%', '<span class="hl">+187%</span>')
                        .replace('YoY +600%', '<span class="hl">YoY +600%</span>')
                        .replace('490+ 雙', '<span class="hl">490+ 雙</span>')}
        </div>
    </div>
    """
    if i < 3:
        exp_html_left += item_html
    else:
        exp_html_right += item_html

html_content = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans+TC:wght@300;400;500;600;700&display=swap" rel="stylesheet">
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: 'Noto Sans TC', sans-serif; background: white; color: #1a1a1a; font-size: 9pt; line-height: 1.4; }}
.resume {{ width: 210mm; min-height: 297mm; padding: 12mm 15mm; margin: 0 auto; background: white; }}
.header {{ display: flex; justify-content: space-between; align-items: flex-start; border-bottom: 2px solid #1a1a1a; padding-bottom: 8px; margin-bottom: 10px; }}
.name {{ font-size: 28pt; font-weight: 700; letter-spacing: -1px; }}
.title {{ font-size: 10pt; color: #666; margin-top: 4px; }}
.contact {{ text-align: right; font-size: 8pt; color: #666; line-height: 1.6; }}
.summary {{ font-size: 9pt; margin-bottom: 12px; line-height: 1.6; }}
.summary b {{ color: #2563eb; }}
.metrics {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px; margin: 12px 0; }}
.metric {{ text-align: center; padding: 8px; background: #f8f9fa; border-radius: 6px; }}
.metric-num {{ font-size: 16pt; font-weight: 700; color: #2563eb; }}
.metric-label {{ font-size: 7pt; color: #666; margin-top: 2px; }}
.section-title {{ font-size: 10pt; font-weight: 700; text-transform: uppercase; letter-spacing: 1px; border-bottom: 1px solid #e5e5e5; padding-bottom: 4px; margin: 14px 0 8px; }}
.two-col {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
.three-col {{ display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 10px; }}
.exp-item {{ margin-bottom: 10px; }}
.exp-header {{ display: flex; justify-content: space-between; align-items: baseline; }}
.exp-company {{ font-weight: 600; font-size: 9pt; }}
.exp-role {{ font-size: 8pt; color: #2563eb; margin-top: 1px; }}
.exp-date {{ font-size: 7.5pt; color: #999; }}
.exp-desc {{ font-size: 8pt; color: #555; margin-top: 3px; line-height: 1.5; }}
.hl {{ color: #2563eb; font-weight: 500; }}
.skill-group-title {{ font-weight: 600; font-size: 8pt; margin-bottom: 5px; color: #333; }}
.skills {{ display: flex; flex-wrap: wrap; gap: 4px; }}
.skill {{ font-size: 7.5pt; background: #f0f0f0; padding: 3px 8px; border-radius: 3px; color: #444; }}
.footer {{ margin-top: 12px; padding-top: 8px; border-top: 1px solid #e5e5e5; font-size: 8pt; color: #666; text-align: center; display: none; }} 
@media print {{ .resume {{ width: 210mm; height: 297mm; padding: 10mm 12mm; }} }}
</style>
</head>
<body>
<div class="resume">
    <div class="header">
        <div>
            <div class="name">{resume_data['name']}</div>
            <div class="title">{resume_data['title']}</div>
        </div>
        <div class="contact">
            {resume_data['contact'][0]}<br>
            {resume_data['contact'][1]}<br>
            {resume_data['contact'][2]}<br>
            {resume_data['contact'][3]}
        </div>
    </div>

    <div class="summary">{resume_data['summary']}</div>

    <div class="metrics">
        {''.join([f'<div class="metric"><div class="metric-num">{m[0]}</div><div class="metric-label">{m[1]}</div></div>' for m in resume_data['metrics']])}
    </div>

    <div class="section-title">工作經歷</div>
    <div style="font-size:7.5pt; color:#666; margin-top:-4px; margin-bottom:6px;">
        近年主要在不同公司負責成長任務與系統重整，多為密集推動期後完成階段目標即離開。
    </div>

    <div class="two-col">
        <div>{exp_html_left}</div>
        <div>{exp_html_right}</div>
    </div>

    <div class="section-title">專業技能</div>
    <div class="three-col">
        <div><div class="skill-group-title">Growth & Strategy</div><div class="skills"><span class="skill">0→1 品牌架構</span><span class="skill">成長策略</span><span class="skill">LTV / CAC</span><span class="skill">漏斗優化</span></div></div>
        <div><div class="skill-group-title">Digital Marketing</div><div class="skills"><span class="skill">FB / Google Ads</span><span class="skill">SEO</span><span class="skill">LINE OA</span><span class="skill">內容行銷</span></div></div>
        <div><div class="skill-group-title">Tools & Platform</div><div class="skills"><span class="skill">GA / GSC</span><span class="skill">Shopline</span><span class="skill">91APP</span><span class="skill">Mixpanel</span></div></div>
    </div>
    <div class="three-col" style="margin-top: 8px;">
        <div><div class="skill-group-title">AI & Automation</div><div class="skills"><span class="skill">AI 內容產製</span><span class="skill">Streamlit Tooling</span><span class="skill">自動化診斷 / 報表</span></div></div>
        <div><div class="skill-group-title">E-commerce & OMO</div><div class="skills"><span class="skill">全通路營運</span><span class="skill">CDP 導入</span><span class="skill">會員經營</span></div></div>
        <div><div class="skill-group-title">Management</div><div class="skills"><span class="skill">跨部門整合</span><span class="skill">SOP 建立</span><span class="skill">專案型推動</span></div></div>
    </div>

    <div class="section-title">代表專案</div>
    <div class="two-col">
        <div class="exp-item">
            <div class="exp-company">🎯 AI 行銷指揮中心</div>
            <div class="exp-desc">
                整合 SEO、廣告、社群、競品監測的一站式工具。日常維運流程自動化後，
                每週例行檢查時間下降約 50%。已應用於保健／美妝／服飾／醫美等品牌。
            </div>
        </div>
        <div class="exp-item">
            <div class="exp-company">📈 可複製的 Growth 流程</div>
            <div class="exp-desc">
                「素材迭代 → 漏斗優化 → 一頁式轉換 → 再行銷矩陣」的成長流程，
                用於服飾、保健食品、3C 等品類，穩定達成 ROAS 4–5。
            </div>
        </div>
    </div>
</div>
</body>
</html>
"""

# 載入 HTML (增加高度以容納新增的經歷)
components.html(html_content, height=1250, scrolling=True)
